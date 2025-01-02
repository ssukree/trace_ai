from difflib import SequenceMatcher
from docx import Document
from docx.shared import RGBColor


def get_differences(prev, curr):
    """Generate the differences between two strings."""
    matcher = SequenceMatcher(None, prev, curr)
    diff_text = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            diff_text.append(("equal", curr[j1:j2]))  # Equal text
        elif tag == 'delete':
            diff_text.append(("delete", prev[i1:i2]))  # Deleted text
        elif tag == 'insert':
            diff_text.append(("insert", curr[j1:j2]))  # Added text
        elif tag == 'replace':
            diff_text.append(("delete", prev[i1:i2]))  # Deleted text
            diff_text.append(("insert", curr[j1:j2]))  # Added text
    
    return diff_text

def create_word_document(cells):
    """Generate a Word document with the comparison of results."""
    doc = Document()

    for index, cell in enumerate(cells):
        doc.add_heading(f'Cell {index}', level=1)
        doc.add_heading(f"Prompt:", level=2)
        doc.add_paragraph(f"{cell['inputs']['prompt']}")
        doc.add_heading(f"Result:",level=2)

        if index == 0:
            # First cell's "Result" as is
            doc.add_paragraph(cell['inputs']['result'])
        else:
            prev_result = cells[index - 1]['inputs']['result']
            curr_result = cell['inputs']['result']
            diff_text = get_differences(prev_result, curr_result)

            # Add the comparison with formatting
            para = doc.add_paragraph()
            for tag, text in diff_text:
                run = para.add_run(text)
                if tag == "delete":
                    run.font.strike = True  # Strike-through for deleted text
                elif tag == "insert":
                    run.font.color.rgb = RGBColor(0, 128, 0)  # Green text
        doc.add_heading(f"Review:",level=2)
        doc.add_paragraph(f"{cell['inputs']['review']}")

    # Save the document
    output_file = 'comparison_report.docx'
    doc.save(output_file)
    return output_file